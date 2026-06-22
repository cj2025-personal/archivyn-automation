"""
Document processing service
Handles PDF, Word, and text file extraction
"""
import pdfplumber
from docx import Document
from typing import Dict, Optional
import requests
import os
import tempfile
from pathlib import Path


class DocumentProcessor:
    """Process various document formats and extract text"""
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def process_document(self, document_url: str) -> Dict:
        """
        Download and process a document (PDF, Word, or text)
        Handles redirects and authentication pages
        
        Returns:
            {
                'content': str,
                'metadata': {...},
                'word_count': int,
                'file_type': str
            }
        """
        try:
            # Download document with redirects and proper headers
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            # Disable SSL verification for self-signed certificates (common on university servers)
            # Suppress urllib3 warnings about unverified HTTPS requests
            import urllib3
            urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
            response = requests.get(document_url, timeout=30, stream=True, headers=headers, allow_redirects=True, verify=False)
            response.raise_for_status()
            
            # Check if we got HTML instead of a document (common with SharePoint/auth pages)
            content_type = response.headers.get('content-type', '').lower()
            if 'text/html' in content_type:
                # Check if it's actually a document by looking at first bytes
                first_bytes = response.content[:512]
                if b'%PDF' in first_bytes[:10] or b'PK\x03\x04' in first_bytes[:10]:
                    # It's actually a document, just mislabeled as HTML
                    pass
                else:
                    # It's really HTML (likely a login/redirect page)
                    # Try Playwright as fallback for auth-required documents
                    print(f"[DocumentProcessor] Direct download returned HTML, trying Playwright fallback for: {document_url[:60]}...")
                    return self._process_document_with_playwright(document_url)
            
            # Determine file type
            file_extension = self._get_file_extension(document_url, content_type)
            
            # If extension couldn't be determined, try to detect from content
            if file_extension == '.pdf' or not file_extension:
                # Check magic numbers to be sure
                first_bytes = response.content[:512]
                if first_bytes[:4] == b'%PDF':
                    file_extension = '.pdf'
                elif first_bytes[:4] == b'PK\x03\x04' and b'word/' in first_bytes[:100]:
                    file_extension = '.docx'
                elif first_bytes[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
                    file_extension = '.doc'
                elif first_bytes[:5] == b'{\\rtf':
                    file_extension = '.rtf'
                elif not file_extension:
                    # Default to PDF if we can't determine
                    file_extension = '.pdf'
            
            # Save to temp file
            temp_file = self._save_temp_file(response.content, file_extension)
            
            try:
                # Process based on file type
                if file_extension == '.pdf':
                    content, metadata = self._process_pdf(temp_file)
                elif file_extension in ['.doc', '.docx']:
                    content, metadata = self._process_word(temp_file)
                elif file_extension == '.txt':
                    content, metadata = self._process_text(temp_file)
                else:
                    raise ValueError(f"Unsupported file type: {file_extension}")
                
                # Debug: Log if content is empty after processing
                if not content or not content.strip():
                    print(f"[DocumentProcessor] WARNING: Processed document {document_url[:60]}... but content is empty. Pages: {metadata.get('pages', 0)}, Has tables: {metadata.get('has_tables', False)}")
                
                return {
                    'content': content if content else '',  # Ensure content is always a string
                    'metadata': metadata,
                    'word_count': len(content.split()) if content else 0,
                    'file_type': file_extension,
                    'file_size': len(response.content),
                    'url': document_url
                }
                
            finally:
                # Clean up temp file
                if os.path.exists(temp_file):
                    os.remove(temp_file)
                    
        except Exception as e:
            error_msg = str(e)
            # If it's an SSL, authentication, or HTML error, try Playwright
            ssl_errors = ['ssl', 'certificate', 'certificate verify failed', 'certificateverificationerror']
            auth_errors = ['authentication', 'html instead of document']
            
            if any(ssl_err in error_msg.lower() for ssl_err in ssl_errors) or any(auth_err in error_msg.lower() for auth_err in auth_errors):
                print(f"[DocumentProcessor] SSL/Auth error detected, trying Playwright fallback for: {document_url[:60]}...")
                try:
                    return self._process_document_with_playwright(document_url)
                except Exception as playwright_error:
                    # If Playwright also fails, raise the original error
                    raise Exception(f"Error processing document {document_url}: {error_msg} (Playwright fallback also failed: {str(playwright_error)})")
            raise Exception(f"Error processing document {document_url}: {error_msg}")
    
    def _process_document_with_playwright(self, document_url: str) -> Dict:
        """
        Download document using Playwright (handles auth-required pages like SharePoint)
        """
        import subprocess
        import json
        import sys
        import os
        import base64
        
        # Find the playwright worker script
        worker_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
            'playwright_subprocess_worker.py'
        )
        
        if not os.path.exists(worker_path):
            raise Exception(f"Playwright worker script not found at {worker_path}")
        
        try:
            task_data = {
                'task': 'download_document',
                'url': document_url
            }
            
            print(f"[DocumentProcessor] Starting document download (timeout: 300s)...")
            result = subprocess.run(
                [sys.executable, worker_path],
                input=json.dumps(task_data),
                capture_output=True,
                text=True,
                timeout=300,  # Increased to 5 minutes for large documents
                encoding='utf-8'
            )
            
            if result.returncode != 0:
                print(f"[DocumentProcessor] Subprocess stderr: {result.stderr[:500]}")
            
            if result.returncode != 0:
                raise Exception(f"Playwright worker failed: {result.stderr}")
            
            # Parse JSON output
            output_lines = result.stdout.strip().split('\n')
            worker_result = None
            for line in output_lines:
                try:
                    worker_result = json.loads(line)
                    break
                except:
                    continue
            
            if not worker_result or worker_result.get('status') != 'success':
                error_msg = worker_result.get('error', 'Unknown error') if worker_result else 'No output'
                raise Exception(f"Playwright download failed: {error_msg}")
            
            # Decode base64 document content
            document_base64 = worker_result.get('document_content', '')
            if not document_base64:
                raise Exception("No document content received from Playwright")
            
            document_content = base64.b64decode(document_base64)
            file_extension = worker_result.get('file_extension', '.pdf')
            
            # Save to temp file and process
            temp_file = self._save_temp_file(document_content, file_extension)
            
            try:
                # Process based on file type
                if file_extension == '.pdf':
                    content, metadata = self._process_pdf(temp_file)
                elif file_extension in ['.doc', '.docx']:
                    content, metadata = self._process_word(temp_file)
                elif file_extension == '.txt':
                    content, metadata = self._process_text(temp_file)
                else:
                    raise ValueError(f"Unsupported file type: {file_extension}")
                
                # Debug: Log if content is empty after processing
                if not content or not content.strip():
                    print(f"[DocumentProcessor] WARNING: Processed document {document_url[:60]}... but content is empty. Pages: {metadata.get('pages', 0)}, Has tables: {metadata.get('has_tables', False)}")
                
                return {
                    'content': content if content else '',  # Ensure content is always a string
                    'metadata': metadata,
                    'word_count': len(content.split()) if content else 0,
                    'file_type': file_extension,
                    'file_size': len(document_content),
                    'url': document_url,
                    'download_method': 'playwright'  # Indicate we used Playwright
                }
            finally:
                # Clean up temp file
                if os.path.exists(temp_file):
                    os.remove(temp_file)
                    
        except subprocess.TimeoutExpired:
            raise Exception("Playwright document download timed out")
        except Exception as e:
            raise Exception(f"Error downloading document with Playwright: {str(e)}")
    
    def _get_file_extension(self, url: str, content_type: str) -> str:
        """Determine file extension from URL or content type"""
        # Check URL first
        url_lower = url.lower()
        if url_lower.endswith('.pdf'):
            return '.pdf'
        elif url_lower.endswith('.docx'):
            return '.docx'
        elif url_lower.endswith('.doc'):
            return '.doc'
        elif url_lower.endswith('.txt'):
            return '.txt'
        elif url_lower.endswith('.rtf'):
            return '.rtf'
        
        # Check content type
        if 'pdf' in content_type:
            return '.pdf'
        elif 'word' in content_type or 'document' in content_type:
            return '.docx'
        elif 'text/plain' in content_type:
            return '.txt'
        
        # Default to PDF if uncertain
        return '.pdf'
    
    def _save_temp_file(self, content: bytes, extension: str) -> str:
        """Save content to temporary file"""
        import uuid
        temp_filename = f"doc_{uuid.uuid4()}{extension}"
        temp_path = os.path.join(self.temp_dir, temp_filename)
        
        with open(temp_path, 'wb') as f:
            f.write(content)
        
        return temp_path
    
    def _process_pdf(self, file_path: str) -> tuple:
        """Extract text from PDF using multiple methods"""
        content_parts = []
        metadata = {
            'pages': 0,
            'has_tables': False,
            'extraction_method': 'pdfplumber'
        }
        
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Method 1: Try pdfplumber's extract_text() (best for text-based PDFs)
                    text = page.extract_text()
                    if text and text.strip():
                        content_parts.append(text)
                    else:
                        # Method 2: Try extract_text with layout preservation
                        text = page.extract_text(layout=True)
                        if text and text.strip():
                            content_parts.append(text)
                        else:
                            # Method 3: Try extracting words and reconstructing
                            words = page.extract_words()
                            if words:
                                # Group words by line (y position)
                                lines = {}
                                for word in words:
                                    y = round(word.get('top', 0))
                                    if y not in lines:
                                        lines[y] = []
                                    # Store both text and x position for sorting
                                    lines[y].append((word.get('x0', 0), word.get('text', '')))
                                
                                # Sort by y position (top to bottom) and x position (left to right)
                                sorted_lines = []
                                for y in sorted(lines.keys(), reverse=True):
                                    # Sort words in line by x position (left to right)
                                    words_in_line = sorted(lines[y], key=lambda w: w[0])
                                    # Extract just the text
                                    line_text = ' '.join([w[1] for w in words_in_line if w[1]])
                                    if line_text.strip():
                                        sorted_lines.append(line_text)
                                
                                if sorted_lines:
                                    content_parts.append('\n'.join(sorted_lines))
                    
                    # Check for tables
                    tables = page.extract_tables()
                    if tables:
                        metadata['has_tables'] = True
                        # Add table data as text
                        for table in tables:
                            if table:
                                table_text = '\n'.join([' | '.join([str(cell) if cell else '' for cell in row]) for row in table if row])
                                if table_text.strip():
                                    content_parts.append(table_text)
        
        except Exception as e:
            # If pdfplumber fails, try PyPDF2 as fallback
            print(f"[DocumentProcessor] pdfplumber extraction failed: {str(e)}, trying PyPDF2 fallback...")
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['pages'] = len(pdf_reader.pages)
                    metadata['extraction_method'] = 'pypdf2'
                    
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text and text.strip():
                            content_parts.append(text)
            except ImportError:
                print("[DocumentProcessor] PyPDF2 not available, skipping fallback")
            except Exception as fallback_error:
                raise Exception(f"Error processing PDF with both pdfplumber and PyPDF2: {str(e)} (PyPDF2: {str(fallback_error)})")
        
        full_content = '\n\n'.join(content_parts)
        
        # If still no content, raise an error with details
        if not full_content.strip():
            raise Exception(f"Could not extract text from PDF. File has {metadata['pages']} pages but no extractable text found. PDF may be image-based or encrypted.")
        
        return full_content, metadata
    
    def _process_word(self, file_path: str) -> tuple:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            
            content_parts = []
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    table_rows.append(row_text)
                if table_rows:
                    content_parts.append('\n'.join(table_rows))
        
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
        
        full_content = '\n\n'.join(content_parts)
        return full_content, metadata
    
    def _process_text(self, file_path: str) -> tuple:
        """Extract text from text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                # Fallback to binary read
                with open(file_path, 'rb') as f:
                    content = f.read().decode('utf-8', errors='ignore')
            
            metadata = {
                'lines': len(content.split('\n'))
            }
            
            return content, metadata
        
        except Exception as e:
            raise Exception(f"Error processing text file: {str(e)}")


# Singleton instance
_document_processor = None

def get_document_processor() -> DocumentProcessor:
    """Get or create document processor instance"""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor



